# VoiceNote UK - FastAPI application entry point
# Serves both the REST API (/transcribe) and the static frontend (/).

# Load .env file automatically so GROQ_API_KEY works on all platforms
from dotenv import load_dotenv
from pathlib import Path

load_dotenv(Path(__file__).with_name(".env"))

import os
import tempfile
import uuid
from typing import List, Optional
from datetime import datetime, date, timedelta
from io import BytesIO

from fastapi import FastAPI, File, HTTPException, UploadFile, Depends, Query, Body, WebSocket
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from sqlalchemy.orm import Session
from pydantic import BaseModel

from models import get_db, User, Note, KeyPoint, ActionItem, Template, UsageStats, Export, create_tables
from auth import (
    get_current_user,
    get_current_user_optional,
    CLERK_ENABLED,
    get_or_create_local_user,
    create_local_auth_response,
)
from templates import get_system_templates, get_template_by_type, get_template_prompt

# ---------------------------------------------------------------------------
# App setup
# ---------------------------------------------------------------------------

app = FastAPI(
    title="VoiceNote UK API",
    description="Transcribe audio in any accent and structure it into clean notes.",
    version="2.0.0",
)

# Initialize database tables on startup
@app.on_event("startup")
async def startup_event():
    """Initialize database tables on startup."""
    try:
        create_tables()
        print("[INFO] Database tables initialized.")
    except Exception as e:
        print(f"[WARNING] Could not initialize database: {e}")
        print("[WARNING] Some features may not work without a database.")

# Allow all origins so the frontend can call the API regardless of hosting setup
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Mount static files (optional)
# app.mount("/static", StaticFiles(directory="static"), name="static")

# Path to the single-file frontend
_FRONTEND_HTML = Path(__file__).parent.parent / "frontend" / "index.html"

# Supported audio extensions
_ALLOWED_EXTENSIONS = {".mp3", ".wav", ".m4a", ".flac", ".ogg"}

# ---------------------------------------------------------------------------
# Pydantic Models
# ---------------------------------------------------------------------------

class NoteCreate(BaseModel):
    title: Optional[str] = None
    transcript: str
    template_type: Optional[str] = "custom"
    audio_file_url: Optional[str] = None
    duration_seconds: Optional[int] = None
    summary: Optional[str] = None
    key_points: Optional[List[str]] = None
    action_items: Optional[List[str]] = None
    tone: Optional[str] = None

class NoteUpdate(BaseModel):
    title: Optional[str] = None
    is_archived: Optional[bool] = None

class TemplateCreate(BaseModel):
    name: str
    template_type: str
    prompt_template: Optional[str] = None
    sections: Optional[dict] = None

class ExportRequest(BaseModel):
    format: str  # 'pdf', 'docx', 'markdown'

class EmailRequest(BaseModel):
    recipients: List[str]
    subject: Optional[str] = None

class LocalLoginRequest(BaseModel):
    email: str
    name: Optional[str] = None


def _analyse_transcript(transcript: str, template_type: str = "custom") -> dict:
    """Import transcript analysis lazily to keep startup fast."""
    from analyse import analyse_transcript

    return analyse_transcript(transcript, template_type)


def _transcribe_audio(audio_path: str) -> str:
    """Import the whisper stack lazily so the server can boot quickly."""
    from transcribe import transcribe_audio

    return transcribe_audio(audio_path)


def _build_note_title(custom_title: Optional[str], template_type: Optional[str] = None) -> str:
    """Return a user-provided title or generate a readable default note title."""
    cleaned_title = (custom_title or "").strip()
    if cleaned_title:
        return cleaned_title

    label_map = {
        "meeting": "Meeting Note",
        "lecture": "Lecture Note",
        "interview": "Interview Note",
        "brainstorm": "Brainstorm Note",
        "custom": "Voice Note",
    }
    base_label = label_map.get((template_type or "custom").lower(), "Voice Note")
    timestamp = datetime.utcnow().strftime("%Y-%m-%d %H:%M")
    return f"{base_label} {timestamp}"


def _should_hide_duplicate_note(note: Note, kept_notes: List[Note]) -> bool:
    """Hide older auto-saved duplicates when a near-identical saved note already exists."""
    transcript = (note.transcript or "").strip()
    if not transcript:
        return False

    for kept in kept_notes:
        same_transcript = (kept.transcript or "").strip() == transcript
        close_in_time = abs((kept.created_at - note.created_at).total_seconds()) <= 180
        if same_transcript and close_in_time:
            return True
    return False


def _dedupe_notes_for_list(notes: List[Note]) -> List[Note]:
    """Return notes ordered newest-first with obvious duplicate autosaves removed."""
    kept_notes: List[Note] = []
    for note in notes:
        if _should_hide_duplicate_note(note, kept_notes):
            continue
        kept_notes.append(note)
    return kept_notes

# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------

@app.get("/", response_class=HTMLResponse, include_in_schema=False)
async def serve_frontend() -> str:
    """Serve the frontend HTML application."""
    return _FRONTEND_HTML.read_text(encoding="utf-8")

@app.get("/healthz", include_in_schema=False)
async def health() -> dict:
    """Quick liveness check."""
    return {"status": "ok"}

# ---------------------------------------------------------------------------
# Authentication Routes
# ---------------------------------------------------------------------------

@app.get("/api/auth/user")
async def get_user(current_user: User = Depends(get_current_user)) -> dict:
    """Get current authenticated user info."""
    return {
        "id": str(current_user.id),
        "email": current_user.email,
        "name": current_user.name,
        "created_at": current_user.created_at.isoformat(),
    }


@app.post("/api/auth/login")
async def local_login(
    login_data: LocalLoginRequest,
    db: Session = Depends(get_db)
) -> dict:
    """Create or fetch a local user and return a signed login token."""
    if CLERK_ENABLED:
        raise HTTPException(status_code=400, detail="Built-in login is disabled when Clerk is configured")

    email = login_data.email.strip()
    if not email or "@" not in email:
        raise HTTPException(status_code=400, detail="A valid email address is required")

    user = get_or_create_local_user(email=email, name=login_data.name, db=db)
    return create_local_auth_response(user)

# ---------------------------------------------------------------------------
# Templates Routes
# ---------------------------------------------------------------------------

@app.get("/api/templates")
async def list_templates(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
) -> dict:
    """Get all available templates (system + user)."""
    # Get system templates
    system_templates = get_system_templates()

    # Get user templates
    user_templates = db.query(Template).filter(
        Template.user_id == current_user.id
    ).all()

    user_template_dicts = [
        {
            "id": str(template.id),
            "name": template.name,
            "template_type": template.template_type,
            "prompt_template": template.prompt_template,
            "sections": template.sections,
            "is_global": False,
            "created_at": template.created_at.isoformat(),
        }
        for template in user_templates
    ]

    return {
        "templates": system_templates + user_template_dicts
    }

@app.post("/api/templates")
async def create_template(
    template_data: TemplateCreate,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
) -> dict:
    """Create a custom template."""
    template = Template(
        user_id=current_user.id,
        name=template_data.name,
        template_type=template_data.template_type,
        prompt_template=template_data.prompt_template,
        sections=template_data.sections,
        is_global=False,
    )

    db.add(template)
    db.commit()
    db.refresh(template)

    return {
        "id": str(template.id),
        "name": template.name,
        "template_type": template.template_type,
        "prompt_template": template.prompt_template,
        "sections": template.sections,
        "is_global": False,
        "created_at": template.created_at.isoformat(),
    }

@app.get("/api/templates/{template_id}")
async def get_template(
    template_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
) -> dict:
    """Get a specific template."""
    # Check if it's a system template first
    if not template_id.startswith("meeting") and not template_id.startswith("lecture") and not template_id.startswith("interview") and not template_id.startswith("brainstorm"):
        try:
            str(uuid.UUID(template_id))
            template = db.query(Template).filter(
                Template.id == template_id,
                Template.user_id == current_user.id
            ).first()

            if not template:
                raise HTTPException(status_code=404, detail="Template not found")

            return {
                "id": str(template.id),
                "name": template.name,
                "template_type": template.template_type,
                "prompt_template": template.prompt_template,
                "sections": template.sections,
                "is_global": False,
                "created_at": template.created_at.isoformat(),
            }
        except ValueError:
            pass

    # Check system templates
    system_template = get_template_by_type(template_id)
    if system_template:
        return system_template

    raise HTTPException(status_code=404, detail="Template not found")

@app.patch("/api/templates/{template_id}")
async def update_template(
    template_id: str,
    template_update: TemplateCreate,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
) -> dict:
    """Update a custom template."""
    try:
        str(uuid.UUID(template_id))
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid template ID")

    template = db.query(Template).filter(
        Template.id == template_id,
        Template.user_id == current_user.id
    ).first()

    if not template:
        raise HTTPException(status_code=404, detail="Template not found")

    # Update fields
    template.name = template_update.name
    template.template_type = template_update.template_type
    template.prompt_template = template_update.prompt_template
    template.sections = template_update.sections

    db.commit()

    return {"message": "Template updated successfully"}

@app.delete("/api/templates/{template_id}")
async def delete_template(
    template_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
) -> dict:
    """Delete a custom template."""
    try:
        str(uuid.UUID(template_id))
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid template ID")

    template = db.query(Template).filter(
        Template.id == template_id,
        Template.user_id == current_user.id
    ).first()

    if not template:
        raise HTTPException(status_code=404, detail="Template not found")

    db.delete(template)
    db.commit()

    return {"message": "Template deleted successfully"}

@app.get("/api/notes")
async def list_notes(
    page: int = Query(1, ge=1),
    limit: int = Query(20, ge=1, le=100),
    search: Optional[str] = None,
    template_type: Optional[str] = None,
    archived: Optional[bool] = None,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
) -> dict:
    """List user's notes with pagination and filtering."""
    query = db.query(Note).filter(Note.user_id == current_user.id)

    # Apply filters
    if search:
        query = query.filter(Note.title.ilike(f"%{search}%") | Note.transcript.ilike(f"%{search}%"))
    if template_type:
        query = query.filter(Note.template_type == template_type)
    if archived is not None:
        query = query.filter(Note.is_archived == archived)

    all_notes = query.order_by(Note.created_at.desc()).all()
    deduped_notes = _dedupe_notes_for_list(all_notes)
    total = len(deduped_notes)
    notes = deduped_notes[(page - 1) * limit : (page - 1) * limit + limit]

    return {
        "notes": [
            {
                "id": str(note.id),
                "title": note.title,
                "summary": note.summary,
                "tone": note.tone,
                "transcript": note.transcript,
                "template_type": note.template_type,
                "audio_file_url": note.audio_file_url,
                "created_at": note.created_at.isoformat(),
                "duration_seconds": note.duration_seconds,
                "word_count": note.word_count,
                "is_archived": note.is_archived,
                "action_items_count": len(note.action_items),
            }
            for note in notes
        ],
        "pagination": {
            "page": page,
            "limit": limit,
            "total": total,
            "pages": (total + limit - 1) // limit,
        }
    }

@app.post("/api/notes")
async def create_note(
    note_data: NoteCreate,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
) -> dict:
    """Create a new note."""
    analysis = {
        "summary": note_data.summary,
        "key_points": note_data.key_points,
        "action_items": note_data.action_items,
        "tone": note_data.tone,
    }
    if not analysis["summary"] and not analysis["key_points"] and not analysis["action_items"] and not analysis["tone"]:
        analysis = _analyse_transcript(note_data.transcript, note_data.template_type or "custom")

    # Create note
    note = Note(
        user_id=current_user.id,
        title=_build_note_title(note_data.title, note_data.template_type),
        transcript=note_data.transcript,
        summary=analysis.get("summary"),
        tone=analysis.get("tone"),
        template_type=note_data.template_type,
        audio_file_url=note_data.audio_file_url,
        duration_seconds=note_data.duration_seconds,
        word_count=len(note_data.transcript.split()),
    )

    db.add(note)
    db.flush()  # Get note ID

    # Add key points
    for i, point in enumerate(analysis.get("key_points", [])):
        key_point = KeyPoint(
            note_id=note.id,
            point=point,
            order_index=i,
        )
        db.add(key_point)

    # Add action items
    for i, item in enumerate(analysis.get("action_items", [])):
        action_item = ActionItem(
            note_id=note.id,
            task=item,
            order_index=i,
        )
        db.add(action_item)

    db.commit()
    db.refresh(note)

    # Track usage
    await _track_note_created(db, current_user.id, note_data.duration_seconds or 0)

    return {
        "id": str(note.id),
        "title": note.title,
        "summary": note.summary,
        "key_points": [kp.point for kp in note.key_points],
        "action_items": [ai.task for ai in note.action_items],
        "tone": note.tone,
        "transcript": note.transcript,
        "template_type": note.template_type,
        "created_at": note.created_at.isoformat(),
    }

@app.get("/api/notes/{note_id}")
async def get_note(
    note_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
) -> dict:
    """Get a specific note."""
    try:
        str(uuid.UUID(note_id))
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid note ID")

    note = db.query(Note).filter(
        Note.id == note_id,
        Note.user_id == current_user.id
    ).first()

    if not note:
        raise HTTPException(status_code=404, detail="Note not found")

    return {
        "id": str(note.id),
        "title": note.title,
        "transcript": note.transcript,
        "summary": note.summary,
        "tone": note.tone,
        "template_type": note.template_type,
        "audio_file_url": note.audio_file_url,
        "created_at": note.created_at.isoformat(),
        "duration_seconds": note.duration_seconds,
        "word_count": note.word_count,
        "is_archived": note.is_archived,
        "key_points": [
            {
                "id": str(kp.id),
                "point": kp.point,
                "order_index": kp.order_index,
                "is_completed": kp.is_completed,
            }
            for kp in note.key_points
        ],
        "action_items": [
            {
                "id": str(ai.id),
                "task": ai.task,
                "order_index": ai.order_index,
                "is_completed": ai.is_completed,
                "assigned_to": ai.assigned_to,
                "due_date": ai.due_date.isoformat() if ai.due_date else None,
            }
            for ai in note.action_items
        ],
    }

@app.patch("/api/notes/{note_id}")
async def update_note(
    note_id: str,
    note_update: NoteUpdate,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
) -> dict:
    """Update a note."""
    try:
        str(uuid.UUID(note_id))
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid note ID")

    note = db.query(Note).filter(
        Note.id == note_id,
        Note.user_id == current_user.id
    ).first()

    if not note:
        raise HTTPException(status_code=404, detail="Note not found")

    # Update fields
    if note_update.title is not None:
        note.title = note_update.title
    if note_update.is_archived is not None:
        note.is_archived = note_update.is_archived

    note.updated_at = datetime.utcnow()
    db.commit()

    return {"message": "Note updated successfully"}

@app.delete("/api/notes/{note_id}")
async def delete_note(
    note_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
) -> dict:
    """Delete a note."""
    try:
        str(uuid.UUID(note_id))
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid note ID")

    note = db.query(Note).filter(
        Note.id == note_id,
        Note.user_id == current_user.id
    ).first()

    if not note:
        raise HTTPException(status_code=404, detail="Note not found")

    db.delete(note)
    db.commit()

    return {"message": "Note deleted successfully"}


@app.get("/api/notes/{note_id}/export")
async def export_note(
    note_id: str,
    format: str = "docx",
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Export a note as a Word document (DOCX)."""
    try:
        str(uuid.UUID(note_id))
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid note ID")

    # In dev mode without auth, just get the note by ID without user filtering
    if not current_user:
        note = db.query(Note).filter(Note.id == note_id).first()
    else:
        note = db.query(Note).filter(
            Note.id == note_id,
            Note.user_id == current_user.id
        ).first()

    if not note:
        raise HTTPException(status_code=404, detail="Note not found")

    if format != "docx":
        raise HTTPException(status_code=400, detail="Only DOCX format is supported")

    # Generate Word document
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Add title
    title = doc.add_heading(note.title or "Voice Note", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add date
    doc.add_paragraph(f"Created: {note.created_at.strftime('%B %d, %Y at %I:%M %p')}")

    # Add summary
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(note.summary or "No summary available")

    # Add tone
    if note.tone:
        doc.add_paragraph(f"Tone: {note.tone}")

    # Add transcript
    doc.add_heading("Transcript", level=1)
    doc.add_paragraph(note.transcript)

    # Add key points
    if note.key_points:
        doc.add_heading("Key Points", level=1)
        for kp in note.key_points:
            doc.add_paragraph(kp.point, style="List Bullet")

    # Add action items
    if note.action_items:
        doc.add_heading("Action Items", level=1)
        for ai in note.action_items:
            doc.add_paragraph(ai.task, style="List Bullet")

    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Track export
    try:
        from datetime import date
        today = date.today()
        stats = db.query(UsageStats).filter(
            UsageStats.user_id == current_user.id,
            UsageStats.date == today
        ).first()
        if stats:
            stats.export_count += 1
            db.commit()
    except Exception:
        pass

    # Return file
    filename = f"{note.title or 'note'}.docx".replace(" ", "_").replace("/", "_")
    headers = {
        "Content-Disposition": f'attachment; filename="{filename}"'
    }
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )

# ---------------------------------------------------------------------------
# Legacy Routes (for backward compatibility)
# ---------------------------------------------------------------------------

@app.post("/transcribe")
async def transcribe_endpoint(
    audio: UploadFile = File(...),
    current_user: Optional[User] = Depends(get_current_user_optional),
) -> dict:
    """
    Accept an audio file, transcribe it with Whisper, and return structured notes.

    Supported formats: mp3, wav, m4a, flac, ogg.

    Returns JSON with:
        transcript   — raw Whisper transcript
        summary      — 2-3 sentence summary
        key_points   — list of bullet-point strings
        action_items — list of task/follow-up strings (may be empty)
        tone         — "formal" | "informal" | "emotional"
    """
    # Validate file extension
    suffix = Path(audio.filename or "").suffix.lower()
    if suffix not in _ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=(
                f"Unsupported format '{suffix}'. "
                f"Allowed: {', '.join(sorted(_ALLOWED_EXTENSIONS))}"
            ),
        )

    # Write upload to a temporary file so faster-whisper can read it
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(await audio.read())
        tmp_path = tmp.name

    try:
        # 1. Transcribe audio → raw text
        transcript = _transcribe_audio(tmp_path)

        # 2. Structure transcript → notes via Groq
        structured = _analyse_transcript(transcript)

        return {"transcript": transcript, **structured}

    finally:
        # Clean up temp file
        Path(tmp_path).unlink(missing_ok=True)

# ---------------------------------------------------------------------------
# Helper Functions
# ---------------------------------------------------------------------------

async def _track_note_created(db: Session, user_id: str, duration_seconds: int):
    """Track note creation in usage stats."""
    try:
        today = date.today()

        # Get or create today's stats
        stats = db.query(UsageStats).filter(
            UsageStats.user_id == user_id,
            UsageStats.date == today
        ).first()

        if not stats:
            stats = UsageStats(
                user_id=user_id,
                date=today,
                notes_created=0,
                transcription_minutes=0,
                export_count=0,
                api_calls=0,
            )
            db.add(stats)

        stats.notes_created += 1
        stats.transcription_minutes += duration_seconds // 60
        db.commit()
    except Exception as e:
        print(f"[WARNING] Could not track usage: {e}")


# ---------------------------------------------------------------------------
# Dev entry point
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    import uvicorn

    port = int(os.environ.get("PORT", 8000))
    uvicorn.run("main:app", host="0.0.0.0", port=port, reload=False)
